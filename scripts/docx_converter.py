"""
Конвертер документов между DOCX и Markdown форматами
"""
import os
import sys
from pathlib import Path
from typing import Optional, Dict, Any
import re

# Добавляем текущую директорию в путь для импорта
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from document_parser import DocumentParser
import markdown2
from html import unescape


class DocxConverter:
    """Конвертер между DOCX и Markdown форматами"""
    
    def __init__(self, documents_dir: str = "documents", 
                 versions_dir: str = "version_history/versions"):
        self.documents_dir = Path(documents_dir)
        self.versions_dir = Path(versions_dir)
        self.parser = DocumentParser(documents_dir)
        self.versions_dir.mkdir(parents=True, exist_ok=True)
    
    def docx_to_markdown(self, docx_path: Path, 
                        include_metadata: bool = True,
                        include_technical: bool = True) -> str:
        """
        Конвертирует DOCX в Markdown
        
        Args:
            docx_path: Путь к DOCX файлу
            include_metadata: Включать ли метаданные в YAML front matter
            include_technical: Включать ли технические данные (статус, даты и т.д.)
        
        Returns:
            Строка с содержимым в формате Markdown
        """
        if not HAS_DOCX:
            raise ImportError("python-docx не установлен. Установите: pip install python-docx")
        
        doc = Document(str(docx_path))
        
        # Извлекаем метаданные из свойств документа
        metadata = {}
        if include_metadata:
            core_props = doc.core_properties
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.created:
                metadata['date'] = core_props.created.strftime('%Y-%m-%d')
            if core_props.modified:
                metadata['modified'] = core_props.modified.strftime('%Y-%m-%d')
        
        # Извлекаем содержимое
        content_lines = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                content_lines.append('')
                continue
            
            # Определяем стиль заголовка
            style_name = paragraph.style.name if paragraph.style else ''
            
            if 'Heading 1' in style_name or paragraph.style.name == 'Title':
                content_lines.append(f'# {text}')
            elif 'Heading 2' in style_name:
                content_lines.append(f'## {text}')
            elif 'Heading 3' in style_name:
                content_lines.append(f'### {text}')
            elif 'Heading 4' in style_name:
                content_lines.append(f'#### {text}')
            elif paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                content_lines.append(f'<center>{text}</center>')
            else:
                content_lines.append(text)
        
        # Обрабатываем таблицы
        for table in doc.tables:
            content_lines.append('')
            # Заголовок таблицы
            if table.rows:
                header_row = table.rows[0]
                headers = [cell.text.strip() for cell in header_row.cells]
                content_lines.append('| ' + ' | '.join(headers) + ' |')
                content_lines.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
                
                # Данные таблицы
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    content_lines.append('| ' + ' | '.join(cells) + ' |')
            content_lines.append('')
        
        content = '\n'.join(content_lines)
        
        # Формируем YAML front matter
        if include_metadata and metadata:
            yaml_lines = ['---']
            for key, value in metadata.items():
                yaml_lines.append(f'{key}: {value}')
            yaml_lines.append('---')
            yaml_lines.append('')
            return '\n'.join(yaml_lines) + content
        
        return content
    
    def markdown_to_docx(self, markdown_content: str, 
                        metadata: Dict[str, Any],
                        output_path: Path,
                        include_metadata: bool = True,
                        include_technical: bool = True) -> Path:
        """
        Конвертирует Markdown в DOCX
        
        Args:
            markdown_content: Содержимое в формате Markdown
            metadata: Метаданные документа
            output_path: Путь для сохранения DOCX файла
            include_metadata: Включать ли метаданные в свойства документа
            include_technical: Включать ли технические данные в документ
        
        Returns:
            Path к созданному DOCX файлу
        """
        if not HAS_DOCX:
            raise ImportError("python-docx не установлен. Установите: pip install python-docx")
        
        doc = Document()
        
        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Устанавливаем метаданные документа
        if include_metadata:
            core_props = doc.core_properties
            if metadata.get('title'):
                core_props.title = metadata['title']
            if metadata.get('organization'):
                core_props.author = metadata.get('organization', '')
            if metadata.get('date'):
                try:
                    from datetime import datetime
                    date_str = metadata['date']
                    if isinstance(date_str, str):
                        # Пробуем разные форматы
                        for fmt in ['%Y-%m-%d', '%d.%m.%Y', '%d/%m/%Y']:
                            try:
                                date_obj = datetime.strptime(date_str.split()[0], fmt)
                                core_props.created = date_obj
                                break
                            except:
                                continue
                except:
                    pass
        
        # Добавляем заголовок
        if metadata.get('title'):
            title = doc.add_heading(metadata['title'], level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем технические данные, если нужно
        if include_technical:
            tech_data = []
            if metadata.get('organization'):
                tech_data.append(f"Организация: {metadata['organization']}")
            if metadata.get('department'):
                tech_data.append(f"Отдел: {metadata['department']}")
            if metadata.get('type'):
                tech_data.append(f"Тип документа: {metadata['type']}")
            if metadata.get('number'):
                tech_data.append(f"Номер: {metadata['number']}")
            if metadata.get('date'):
                tech_data.append(f"Дата: {metadata['date']}")
            if metadata.get('status'):
                tech_data.append(f"Статус: {metadata['status']}")
            
            if tech_data:
                tech_para = doc.add_paragraph('\n'.join(tech_data))
                tech_para.style = 'List Bullet'
                doc.add_paragraph('')  # Пустая строка
        
        # Конвертируем Markdown в HTML для более точной обработки
        html_content = markdown2.markdown(
            markdown_content,
            extras=['fenced-code-blocks', 'tables', 'header-ids', 'break-on-newline']
        )
        
        # Используем простой парсинг HTML, который лучше сохраняет форматирование
        self._simple_html_parse(html_content, doc)
        
        # Сохраняем документ
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        return output_path
    
    def _html_to_docx(self, html_content: str, doc: Document):
        """
        Конвертирует HTML в элементы DOCX
        
        Args:
            html_content: HTML содержимое
            doc: Объект Document для добавления элементов
        """
        from xml.etree import ElementTree as ET
        import re
        
        # Парсим HTML
        try:
            # Очищаем HTML от некорректных тегов
            html_content = re.sub(r'<br\s*/?>', '\n', html_content)
            root = ET.fromstring(f'<root>{html_content}</root>')
        except:
            # Если не удалось распарсить как XML, используем простой парсинг
            self._simple_html_parse(html_content, doc)
            return
        
        current_para = None
        list_level = 0
        list_style = None
        
        def process_element(elem, para=None, in_list=False):
            nonlocal current_para, list_level, list_style
            
            tag = elem.tag.lower()
            text = elem.text or ''
            tail = elem.tail or ''
            
            # Заголовки
            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag[1])
                heading_text = self._extract_text(elem)
                if heading_text.strip():
                    doc.add_heading(heading_text, level=min(level, 4))
                current_para = None
            
            # Параграфы
            elif tag == 'p':
                para_text = self._extract_text(elem)
                if para_text.strip():
                    current_para = doc.add_paragraph()
                    self._add_formatted_text(current_para, para_text)
                else:
                    doc.add_paragraph('')
            
            # Списки
            elif tag in ['ul', 'ol']:
                list_level += 1
                list_style = 'List Bullet' if tag == 'ul' else 'List Number'
                for item in elem.findall('.//li'):
                    item_text = self._extract_text(item)
                    if item_text.strip():
                        para = doc.add_paragraph(item_text, style=list_style)
                        self._add_formatted_text(para, item_text)
                list_level -= 1
                list_style = None
            
            # Элементы списка
            elif tag == 'li':
                if not in_list:
                    item_text = self._extract_text(elem)
                    if item_text.strip():
                        para = doc.add_paragraph(item_text, style=list_style or 'List Bullet')
                        self._add_formatted_text(para, item_text)
            
            # Таблицы
            elif tag == 'table':
                self._process_table(elem, doc)
            
            # Обычный текст
            elif text.strip() or tail.strip():
                if not current_para:
                    current_para = doc.add_paragraph()
                if text.strip():
                    self._add_formatted_text(current_para, text)
                if tail.strip():
                    self._add_formatted_text(current_para, tail)
            
            # Рекурсивно обрабатываем дочерние элементы
            for child in elem:
                process_element(child, current_para, in_list=(tag in ['ul', 'ol']))
        
        # Обрабатываем корневой элемент
        for child in root:
            process_element(child)
    
    def _simple_html_parse(self, html_content: str, doc: Document):
        """Простой парсинг HTML через регулярные выражения с сохранением форматирования"""
        import re
        
        # Обрабатываем таблицы отдельно
        tables = []
        table_pattern = r'<table[^>]*>(.*?)</table>'
        for match in re.finditer(table_pattern, html_content, flags=re.DOTALL | re.IGNORECASE):
            tables.append((match.start(), match.end(), match.group(0)))
            html_content = html_content[:match.start()] + f'__TABLE_{len(tables)-1}__' + html_content[match.end():]
        
        # Обрабатываем списки с сохранением структуры
        html_content = re.sub(r'<ul[^>]*>', '\n__UL_START__\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'</ul>', '\n__UL_END__\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'<ol[^>]*>', '\n__OL_START__\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'</ol>', '\n__OL_END__\n', html_content, flags=re.IGNORECASE)
        
        # Заголовки
        html_content = re.sub(r'<h1[^>]*>(.*?)</h1>', r'__H1__\1__/H1__', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<h2[^>]*>(.*?)</h2>', r'__H2__\1__/H2__', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<h3[^>]*>(.*?)</h3>', r'__H3__\1__/H3__', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<h4[^>]*>(.*?)</h4>', r'__H4__\1__/H4__', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Параграфы
        html_content = re.sub(r'<p[^>]*>(.*?)</p>', r'__P__\1__/P__', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Элементы списков
        html_content = re.sub(r'<li[^>]*>(.*?)</li>', r'__LI__\1__/LI__', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Удаляем оставшиеся HTML теги, но сохраняем форматирование
        # Жирный и курсив обрабатываем отдельно
        html_content = re.sub(r'<strong[^>]*>(.*?)</strong>', r'__BOLD__\1__/BOLD__', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<b[^>]*>(.*?)</b>', r'__BOLD__\1__/BOLD__', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<em[^>]*>(.*?)</em>', r'__ITALIC__\1__/ITALIC__', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<i[^>]*>(.*?)</i>', r'__ITALIC__\1__/ITALIC__', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Удаляем остальные теги
        html_content = re.sub(r'<[^>]+>', '', html_content)
        html_content = unescape(html_content)
        
        # Восстанавливаем таблицы
        for idx, (start, end, table_html) in enumerate(tables):
            html_content = html_content.replace(f'__TABLE_{idx}__', table_html)
        
        # Парсим построчно
        lines = html_content.split('\n')
        i = 0
        in_list = False
        list_style = None
        last_was_empty = False  # Отслеживаем последовательные пустые строки
        
        while i < len(lines):
            line = lines[i]
            
            # Таблицы
            if '__TABLE_' in line or '<table' in line.lower():
                # Ищем полную таблицу
                table_text = line
                j = i + 1
                while j < len(lines) and ('__TABLE_' in lines[j] or '</table>' not in lines[j]):
                    table_text += '\n' + lines[j]
                    j += 1
                if '</table>' in lines[j]:
                    table_text += '\n' + lines[j]
                
                self._parse_table_html(table_text, doc)
                i = j + 1
                last_was_empty = False
                continue
            
            # Заголовки
            if '__H1__' in line:
                text = re.sub(r'__H1__|__/H1__', '', line)
                text = re.sub(r'__(P|/P|LI|/LI|UL_START|UL_END|OL_START|OL_END)__', '', text).strip()
                if text:
                    doc.add_heading(text, level=1)
                    last_was_empty = False
            elif '__H2__' in line:
                text = re.sub(r'__H2__|__/H2__', '', line)
                text = re.sub(r'__(P|/P|LI|/LI|UL_START|UL_END|OL_START|OL_END)__', '', text).strip()
                if text:
                    doc.add_heading(text, level=2)
                    last_was_empty = False
            elif '__H3__' in line:
                text = re.sub(r'__H3__|__/H3__', '', line)
                text = re.sub(r'__(P|/P|LI|/LI|UL_START|UL_END|OL_START|OL_END)__', '', text).strip()
                if text:
                    doc.add_heading(text, level=3)
                    last_was_empty = False
            elif '__H4__' in line:
                text = re.sub(r'__H4__|__/H4__', '', line)
                text = re.sub(r'__(P|/P|LI|/LI|UL_START|UL_END|OL_START|OL_END)__', '', text).strip()
                if text:
                    doc.add_heading(text, level=4)
                    last_was_empty = False
            # Списки
            elif '__UL_START__' in line:
                in_list = True
                list_style = 'List Bullet'
                last_was_empty = False
            elif '__UL_END__' in line or '__OL_END__' in line:
                in_list = False
                list_style = None
                last_was_empty = False
            elif '__OL_START__' in line:
                in_list = True
                list_style = 'List Number'
                last_was_empty = False
            elif '__LI__' in line:
                # Извлекаем текст, удаляя маркеры LI
                text = re.sub(r'__LI__|__/LI__', '', line)
                # Удаляем все остальные служебные маркеры
                text = re.sub(r'__(P|/P|H[1-6]|/H[1-6]|UL_START|UL_END|OL_START|OL_END)__', '', text).strip()
                if text:
                    para = doc.add_paragraph(text, style=list_style or 'List Bullet')
                    self._add_formatted_text_from_html(para, text)
                    last_was_empty = False
            # Параграфы
            elif '__P__' in line:
                # Извлекаем текст, удаляя маркеры P
                text = re.sub(r'__P__|__/P__', '', line)
                # Удаляем все остальные служебные маркеры
                text = re.sub(r'__(LI|/LI|H[1-6]|/H[1-6]|UL_START|UL_END|OL_START|OL_END)__', '', text).strip()
                if text:
                    para = doc.add_paragraph()
                    self._add_formatted_text_from_html(para, text)
                    last_was_empty = False
            # Обычный текст
            elif line.strip() and not line.startswith('__'):
                para = doc.add_paragraph()
                self._add_formatted_text_from_html(para, line.strip())
                last_was_empty = False
            # Пустые строки - добавляем только одну, если предыдущая не была пустой
            elif not line.strip():
                # Пропускаем пустые строки, если предыдущая тоже была пустой
                # или если мы в списке
                if not last_was_empty and not in_list:
                    # Добавляем только одну пустую строку для разделения блоков
                    last_was_empty = True
                # Иначе просто пропускаем
            
            i += 1
    
    def _add_formatted_text_from_html(self, para, text: str):
        """Добавляет форматированный текст из HTML-подобных маркеров"""
        import re
        
        para.clear()
        
        # Сначала удаляем все служебные маркеры (кроме форматирования)
        # Удаляем маркеры параграфов, списков, заголовков
        text = re.sub(r'__(P|/P|LI|/LI|H[1-6]|/H[1-6]|UL_START|UL_END|OL_START|OL_END)__', '', text)
        
        # Сначала обрабатываем правильно оформленные маркеры форматирования
        # Ищем пары __BOLD__...__/BOLD__ и __ITALIC__...__/ITALIC__
        # Используем нежадный поиск, чтобы правильно обработать вложенные маркеры
        # Важно: проверяем, что это действительно правильная пара (не поврежденная)
        parts = re.split(r'(__BOLD__.*?__/BOLD__|__ITALIC__.*?__/ITALIC__)', text, flags=re.DOTALL)
        
        for part in parts:
            if not part:
                continue
            
            # Обрабатываем правильно оформленные пары маркеров
            # Проверяем, что это действительно правильная пара (не множественные маркеры подряд)
            if part.startswith('__BOLD__') and part.endswith('__/BOLD__'):
                content = part[8:-9]  # Убираем __BOLD__ и __/BOLD__
                # Проверяем, что внутри нет множественных маркеров (поврежденный случай)
                # Если внутри есть другие маркеры или только подчеркивания - это поврежденный случай
                if '__BOLD__' in content or '__ITALIC__' in content or '__/BOLD__' in content or '__/ITALIC__' in content:
                    # Это поврежденный маркер, удаляем его полностью
                    clean_text = re.sub(r'__+[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?__+', '', part)
                    clean_text = re.sub(r'[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?__+', '', clean_text)
                    clean_text = re.sub(r'__+[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?', '', clean_text)
                    # Удаляем одиночные подчеркивания в начале/конце (остатки от маркеров)
                    clean_text = re.sub(r'^_+\s*|\s*_+$', '', clean_text)
                    clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                    if clean_text:
                        para.add_run(clean_text)
                elif content.strip() and not content.strip().startswith('_') and len(content.strip()) > 1:
                    # Рекурсивно обрабатываем вложенное форматирование
                    self._add_formatted_text_to_para(para, content, force_bold=True)
                else:
                    # Пустой или только подчеркивания - удаляем маркеры
                    clean_text = re.sub(r'__+[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?__+', '', part)
                    clean_text = re.sub(r'[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?__+', '', clean_text)
                    clean_text = re.sub(r'__+[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?', '', clean_text)
                    clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                    if clean_text:
                        para.add_run(clean_text)
            elif part.startswith('__ITALIC__') and part.endswith('__/ITALIC__'):
                content = part[10:-11]  # Убираем __ITALIC__ и __/ITALIC__
                # Проверяем, что внутри нет множественных маркеров (поврежденный случай)
                if '__BOLD__' in content or '__ITALIC__' in content or '__/BOLD__' in content or '__/ITALIC__' in content:
                    # Это поврежденный маркер, удаляем его полностью
                    clean_text = re.sub(r'__+[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?__+', '', part)
                    clean_text = re.sub(r'[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?__+', '', clean_text)
                    clean_text = re.sub(r'__+[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?', '', clean_text)
                    # Удаляем одиночные подчеркивания в начале/конце (остатки от маркеров)
                    clean_text = re.sub(r'^_+\s*|\s*_+$', '', clean_text)
                    clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                    if clean_text:
                        para.add_run(clean_text)
                elif content.strip() and not content.strip().startswith('_') and len(content.strip()) > 1:
                    # Рекурсивно обрабатываем вложенное форматирование
                    self._add_formatted_text_to_para(para, content, force_italic=True)
                else:
                    # Пустой или только подчеркивания - удаляем маркеры
                    clean_text = re.sub(r'__+[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?__+', '', part)
                    clean_text = re.sub(r'[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?__+', '', clean_text)
                    clean_text = re.sub(r'__+[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?', '', clean_text)
                    clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                    if clean_text:
                        para.add_run(clean_text)
            else:
                # Для обычного текста удаляем все поврежденные маркеры
                # Удаляем неполные/поврежденные маркеры форматирования
                clean_text = re.sub(r'__+[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?__+', '', part)
                clean_text = re.sub(r'[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?__+', '', clean_text)
                clean_text = re.sub(r'__+[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?', '', clean_text)
                # Удаляем одиночные подчеркивания в начале/конце (остатки от маркеров)
                clean_text = re.sub(r'^_+\s*|\s*_+$', '', clean_text)
                # Удаляем множественные пробелы и пробелы в начале/конце
                clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                if clean_text:
                    para.add_run(clean_text)
    
    def _add_formatted_text_to_para(self, para, text: str, force_bold=False, force_italic=False):
        """Вспомогательный метод для добавления форматированного текста с принудительным форматированием"""
        import re
        
        # Обрабатываем вложенное форматирование
        parts = re.split(r'(__BOLD__.*?__/BOLD__|__ITALIC__.*?__/ITALIC__)', text, flags=re.DOTALL)
        
        for part in parts:
            if not part:
                continue
            
            if part.startswith('__BOLD__') and part.endswith('__/BOLD__'):
                content = part[8:-9]
                if content.strip():
                    self._add_formatted_text_to_para(para, content, force_bold=True, force_italic=force_italic)
            elif part.startswith('__ITALIC__') and part.endswith('__/ITALIC__'):
                content = part[10:-11]
                if content.strip():
                    self._add_formatted_text_to_para(para, content, force_bold=force_bold, force_italic=True)
            else:
                # Удаляем поврежденные маркеры
                clean_text = re.sub(r'__+[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?__+', '', part)
                clean_text = re.sub(r'[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?__+', '', clean_text)
                clean_text = re.sub(r'__+[A-Z_/]*?(BOLD|ITALIC|/BOLD|/ITALIC)[A-Z_/]*?', '', clean_text)
                clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                if clean_text:
                    run = para.add_run(clean_text)
                    if force_bold:
                        run.bold = True
                    if force_italic:
                        run.italic = True
    
    def _parse_table_html(self, table_html: str, doc: Document):
        """Парсит HTML таблицу и добавляет в документ"""
        import re
        
        # Извлекаем строки
        rows = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html, flags=re.DOTALL | re.IGNORECASE)
        if not rows:
            return
        
        # Определяем количество столбцов
        max_cols = 0
        table_data = []
        for row in rows:
            cells = re.findall(r'<(td|th)[^>]*>(.*?)</(td|th)>', row, flags=re.DOTALL | re.IGNORECASE)
            cell_texts = [unescape(cell[1].strip()) for cell in cells]
            table_data.append((cell_texts, 'th' in row.lower()))
            max_cols = max(max_cols, len(cell_texts))
        
        if max_cols == 0:
            return
        
        # Создаем таблицу
        table = doc.add_table(rows=1, cols=max_cols)
        table.style = 'Light Grid Accent 1'
        
        for row_idx, (cells, is_header) in enumerate(table_data):
            if row_idx > 0:
                table.add_row()
            
            for col_idx, cell_text in enumerate(cells[:max_cols]):
                table.rows[row_idx].cells[col_idx].text = cell_text
                if is_header or row_idx == 0:
                    para = table.rows[row_idx].cells[col_idx].paragraphs[0]
                    for run in para.runs:
                        run.bold = True
    
    def _extract_text(self, elem) -> str:
        """Извлекает текст из HTML элемента"""
        text = elem.text or ''
        for child in elem:
            text += self._extract_text(child)
        text += elem.tail or ''
        return unescape(text).strip()
    
    def _add_formatted_text(self, para, text: str):
        """Добавляет форматированный текст в параграф"""
        import re
        
        # Очищаем параграф от текста по умолчанию
        para.clear()
        
        # Обрабатываем жирный и курсив
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        
        for part in parts:
            if not part:
                continue
            
            # Жирный текст
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            # Курсив
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                run = para.add_run(part[1:-1])
                run.italic = True
            # Обычный текст
            else:
                para.add_run(part)
    
    def _process_table(self, table_elem, doc: Document):
        """Обрабатывает HTML таблицу"""
        rows = table_elem.findall('.//tr')
        if not rows:
            return
        
        # Определяем количество столбцов
        max_cols = 0
        for row in rows:
            cols = len(row.findall('.//td') + row.findall('.//th'))
            max_cols = max(max_cols, cols)
        
        if max_cols == 0:
            return
        
        # Создаем таблицу
        table = doc.add_table(rows=1, cols=max_cols)
        table.style = 'Light Grid Accent 1'
        
        for row_idx, row_elem in enumerate(rows):
            if row_idx > 0:
                table.add_row()
            
            cells = row_elem.findall('.//td') + row_elem.findall('.//th')
            for col_idx, cell_elem in enumerate(cells[:max_cols]):
                cell_text = self._extract_text(cell_elem)
                table.rows[row_idx].cells[col_idx].text = cell_text
                
                # Жирный для заголовков
                if cell_elem.tag.lower() == 'th':
                    para = table.rows[row_idx].cells[col_idx].paragraphs[0]
                    for run in para.runs:
                        run.bold = True
    
    def save_docx_version(self, docx_path: Path, doc_relative_path: str, 
                         author: str, comment: Optional[str] = None) -> Path:
        """
        Сохраняет DOCX файл как версию документа
        
        Args:
            docx_path: Путь к загруженному DOCX файлу
            doc_relative_path: Относительный путь к документу в documents/
            author: Автор загрузки
            comment: Комментарий к версии
        
        Returns:
            Path к сохраненной версии
        """
        from datetime import datetime
        
        # Создаем структуру директорий для версий
        doc_path = self.documents_dir / doc_relative_path
        rel_path = doc_path.relative_to(self.documents_dir)
        version_dir = self.versions_dir / rel_path.parent
        version_dir.mkdir(parents=True, exist_ok=True)
        
        # Имя файла версии с временной меткой
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        version_filename = f"{doc_path.stem}_{timestamp}.docx"
        version_path = version_dir / version_filename
        
        # Копируем DOCX файл
        import shutil
        shutil.copy2(docx_path, version_path)
        
        return version_path

